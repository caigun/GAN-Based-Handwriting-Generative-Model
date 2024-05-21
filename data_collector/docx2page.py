import os
from docx import Document
from PIL import Image
import re


def extract_title_and_images(docx_path):
    document = Document(docx_path)
    element_count = len(document.paragraphs)
    print(element_count)
    print(len(document.inline_shapes))
    images = []
    current_title = None
    title_images = {}
    image_index = 0
    title_index = 0
    for i in range(element_count):
        paragraph = document.paragraphs[i]

        if paragraph.text.strip() == "":
            if len(paragraph.runs) > 0:
                if len(paragraph.runs) > 1:
                    print("exist more than one element in one line", paragraph.runs)
                for run in paragraph.runs:
                    # print(image_index)
                    if run.text != "":  # if there's a blank space
                        print(paragraph.runs)
                        print("!!!!!!blank space!!!!!!!")
                        continue

                    image_data = document.inline_shapes[image_index]._inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_bytes = document.part.related_parts[image_data].blob
                    images.append(image_bytes)
                    image_index += 1
            else:
                print("----------empty line-----------")  # 纯空行，连空格都没有
                pass
        else:  # for text
            # print(paragraph.text)
            if images:
                title_images[current_title] = images
                images = []
            current_title = title_index  # paragraph.text.strip()
            title_index += 1

    if images:
        title_images[current_title] = images
    print(list(title_images.keys()))
    return title_images


def process_images(title_images, output_folder):
    for title, images in title_images.items():

        # one dir per person
        # title_folder = os.path.join(output_folder, title)
        # os.makedirs(title_folder, exist_ok=True)

        for index, image_bytes in enumerate(images, start=1):
            # output_folder --> title_folder.
            new_image_path = os.path.join(
                output_folder, f"{title}_{index}.png")
            with open(new_image_path, "wb") as img_file:
                img_file.write(image_bytes)


def main(docx_path, output_folder):
    title_images = extract_title_and_images(docx_path)
    process_images(title_images, output_folder)


data_dir = 'data'
new_dataset_dir = os.path.join(data_dir, 'new_dataset')
idx = 1
for file in os.listdir(new_dataset_dir):
    if file.endswith(".docx"):
        docx_path = os.path.join(new_dataset_dir, file)
        output_folder = os.path.join(
            new_dataset_dir, "_page_img_" + str(idx))
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
        main(docx_path, output_folder)
        idx += 1
